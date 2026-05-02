from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Rapport_PFA_Plateforme_IA_EMSI.docx"

GREEN = RGBColor(31, 122, 77)
DARK_GREEN = RGBColor(20, 92, 56)
LIGHT_GREEN = "E8F5EE"
HEADER_BG = "1F7A4D"
GRAY_BG = "F4F6F5"
WHITE = RGBColor(255, 255, 255)
DARK_GRAY = RGBColor(64, 64, 64)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C7D5CE", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char2)
    paragraph.add_run("Cliquez droit puis 'Mettre a jour le champ' si Word ne l'actualise pas automatiquement.")
    run2 = paragraph.add_run()
    run2._r.append(fld_char3)


def mark_update_fields(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 22, GREEN),
        ("Heading 1", 17, GREEN),
        ("Heading 2", 14, GREEN),
        ("Heading 3", 12, DARK_GREEN),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 9)
        st.paragraph_format.space_after = Pt(5)

    if "CaptionCustom" not in styles:
        cap = styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Arial"
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.font.color.rgb = DARK_GRAY
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(9)


def header_footer(doc):
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(4)
    r = hp.add_run("Rapport PFA - Plateforme IA de Recommandation | EMSI 3IIR G3 - 2025/2026")
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.color.rgb = DARK_GRAY
    hp._p.get_or_add_pPr().append(bottom_border())

    fp = section.footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp._p.get_or_add_pPr().append(top_border())
    run = fp.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    add_page_field(fp, "PAGE")
    fp.add_run(" / ")
    add_page_field(fp, "NUMPAGES")


def bottom_border():
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1F7A4D")
    p_bdr.append(bottom)
    return p_bdr


def top_border():
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "2")
    top.set(qn("w:color"), "1F7A4D")
    p_bdr.append(top)
    return p_bdr


def add_page_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, end])


def p(doc, text="", bold=False, italic=False, color=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.space_after = Pt(6)
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = color
    return par


def add_mixed(doc, runs, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    par = doc.add_paragraph()
    par.alignment = align
    for item in runs:
        run = par.add_run(item["text"])
        run.bold = item.get("bold", False)
        run.italic = item.get("italic", False)
        run.font.name = "Arial"
        run.font.size = Pt(item.get("size", 10.5))
        if item.get("color"):
            run.font.color.rgb = item["color"]
    return par


def h(doc, text, level=1):
    par = doc.add_heading(text, level=level)
    if level == 1:
        par._p.get_or_add_pPr().append(bottom_border())
    return par


def bullet(doc, text, level=0):
    par = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    par.paragraph_format.space_after = Pt(2.5)
    run = par.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def numbered(doc, text):
    par = doc.add_paragraph(style="List Number")
    par.paragraph_format.space_after = Pt(2.5)
    run = par.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def caption(doc, text):
    doc.add_paragraph(text, style="CaptionCustom")


def table(doc, headers, rows, widths=None):
    cols = len(headers)
    tbl = doc.add_table(rows=1, cols=cols)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl.autofit = True
    set_table_borders(tbl)
    tr_pr = tbl.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    hdr = tbl.rows[0].cells
    for i, head in enumerate(headers):
        set_cell_shading(hdr[i], HEADER_BG)
        set_cell_margins(hdr[i])
        set_cell_text(hdr[i], head, bold=True, color=WHITE, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci, value in enumerate(row):
            set_cell_shading(cells[ci], "FFFFFF" if ri % 2 == 0 else GRAY_BG)
            set_cell_margins(cells[ci])
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[ci], value, size=8.9, align=align)
    if widths:
        for row in tbl.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    return tbl


def figure_flow(doc, title, labels):
    tbl = doc.add_table(rows=1, cols=len(labels))
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(tbl, "8BB89E", "8")
    tr_pr = tbl.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, label in enumerate(labels):
        cell = tbl.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GREEN if i % 2 == 0 else "FFFFFF")
        set_cell_margins(cell, top=150, bottom=150, start=130, end=130)
        set_cell_text(cell, label, bold=True, color=DARK_GREEN, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    caption(doc, title)


def add_cover(doc):
    for _ in range(3):
        p(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    par = p(doc, "EMSI - Ecole Marocaine des Sciences de l'Ingenieur", bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)
    par.runs[0].font.size = Pt(16)
    p(doc, "Filiere : Ingenierie Informatique - 3eme annee (3IIR - G3)", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Annee universitaire 2025 - 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()._p.get_or_add_pPr().append(bottom_border())
    for _ in range(3):
        p(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    title = p(doc, "RAPPORT DE PROJET DE FIN D'ANNEE", bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.runs[0].font.size = Pt(20)
    main = p(doc, "Plateforme Intelligente de Recommandation de Stages et d'Emplois par Intelligence Artificielle", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    main.runs[0].font.size = Pt(18)
    p(doc, "Conception, implementation et evaluation d'un moteur NLP hybride, d'une architecture SaaS Django/React et d'un module de scraping oriente marche marocain.", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        p(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    table(doc, ["Element", "Information"], [
        ["Realise par", "Ayman Ould-Naghouch et El Mehdi Ribahi"],
        ["Encadrant", "M. Gamal Mohamed"],
        ["Etablissement", "EMSI - Casablanca, Maroc"],
        ["Projet", "Plateforme IA de recommandation de stages et d'emplois"],
        ["Annee", "2025 - 2026"],
    ], [4, 11])
    p(doc, "Casablanca, Maroc | 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_intro(doc):
    h(doc, "Introduction generale", 1)
    p(doc, "Dans un contexte mondial ou la transformation numerique redefinit les pratiques de recrutement, la recherche de stage ou d'emploi chez les etudiants demeure encore largement manuelle et peu efficiente. Les candidats font face a un volume croissant d'offres heterogenes, a une faible personnalisation des resultats de recherche, et a une difficulte reelle d'evaluer objectivement la pertinence d'une opportunite par rapport a leur profil.")
    p(doc, "Au Maroc, ce phenomene est particulierement marque. Le marche de l'emploi pour les jeunes ingenieurs est competitif, et les plateformes disponibles proposent principalement une approche de job board passif basee sur des filtres mots-cles, sans analyse semantique ni personnalisation intelligente.")
    p(doc, "Le projet PFA developpe une plateforme complete qui comprend l'upload et le parsing de CV, l'extraction de competences, la recommandation explicable, l'assistance IA a la candidature, le pilotage entreprise et la veille marche via scraping.")
    add_mixed(doc, [{"text": "Objectif principal : ", "bold": True, "color": GREEN}, {"text": "concevoir et implementer une plateforme intelligente de recommandation de stages et d'emplois, orientee marche marocain, combinant NLP, IA generative et architecture SaaS moderne."}])


def add_conceptual(doc):
    h(doc, "Chapitre 1 - Cadre conceptuel : du e-recrutement a l'IA de recommandation", 1)
    h(doc, "1.1 Job Recommender Systems", 2)
    p(doc, "Les Job Recommender Systems designent les systemes capables d'analyser automatiquement le profil d'un candidat et de lui proposer de facon proactive les opportunites professionnelles les plus pertinentes.")
    table(doc, ["Phase", "Periode", "Caracteristiques", "Position du projet"], [
        ["Job boards passifs", "2000-2010", "Recherche par mots-cles et tri manuel", "Contexte initial"],
        ["Matching semi-automatique", "2010-2018", "Filtres avances sans comprehension semantique profonde", "Limite observee"],
        ["Recommandation proactive", "2018-aujourd'hui", "NLP, profilage, scoring, explication", "Cible du PFA"],
    ], [3, 2.5, 6, 4])
    caption(doc, "Figure 1 - Evolution chronologique des approches de recommandation d'emploi")
    h(doc, "1.2 Apprentissage automatique", 2)
    p(doc, "L'apprentissage automatique permet a un systeme d'apprendre des regularites a partir de donnees. Pour la recommandation, il sert a modeliser la pertinence entre un profil candidat, ses preferences et les caracteristiques des offres.")
    bullet(doc, "Apprentissage supervise : modeles entraines sur des exemples etiquetes.")
    bullet(doc, "Apprentissage non supervise : detection de similarites, clusters ou structures cachees.")
    bullet(doc, "Apprentissage semi-supervise : utile lorsque les donnees annotees sont rares.")
    h(doc, "1.3 Resume Parsing", 2)
    p(doc, "Le resume parsing est le processus d'extraction d'informations structurees depuis un CV : competences, formation, experience, projets et preferences. Dans ce projet, il combine extraction de texte, heuristiques et, lorsque disponible, IA generative.")
    h(doc, "1.4 TF-IDF et similarite cosinus", 2)
    p(doc, "TF-IDF pondere les termes selon leur importance dans un document et leur rarete dans le corpus. La similarite cosinus compare ensuite les vecteurs du CV et des offres afin d'obtenir un score independant de la longueur des textes.")
    h(doc, "1.5 Architecture hybride", 2)
    table(doc, ["Approche", "Principe", "Avantage", "Statut"], [
        ["Content-Based", "Matching selon contenu profil/offre", "Gerable au demarrage", "Implemente"],
        ["Collaborative Filtering", "Matching selon comportements similaires", "Pertinent avec historique", "Perspective"],
        ["Hybride", "Fusion contenu + comportement", "Etat de l'art mature", "Roadmap"],
    ], [3.2, 5, 4, 2.5])
    caption(doc, "Tableau 1 - Comparaison des architectures de recommandation")


def add_methodology(doc):
    h(doc, "Chapitre 2 - Cadre methodologique et revue de litterature", 1)
    h(doc, "2.1 Criteres de selection des sources", 2)
    bullet(doc, "Sources Q1/Q2, IEEE/ACM ou publications reconnues.")
    bullet(doc, "Pertinence thematique : recommandation d'emploi, NLP, parsing de CV, systemes SaaS.")
    bullet(doc, "Periode privilegiee : 2012-2024, avec references fondamentales lorsque necessaire.")
    bullet(doc, "Accessibilite et tracabilite des sources pour verification.")
    h(doc, "2.2 Tableau bibliographique synthetique", 2)
    table(doc, ["Ref.", "Auteurs", "Apport", "Quartile", "Support"], [
        ["[1]", "Isinkaye et al., 2015", "Principes, methodes et evaluation des systemes de recommandation", "Q1", "Egyptian Informatics Journal"],
        ["[2]", "Siting et al., 2012", "Survey des Job Recommender Systems", "Q2", "ICCSE"],
        ["[3]", "Al-Otaibi & Ykhlef, 2012", "Taxonomie des systemes RH de recommandation", "Q2", "IJPS"],
        ["[4]", "Yi et al., 2020", "Baseline TF-IDF + cosine pour emploi", "Q2", "ICMLC"],
        ["[5]", "Kops et al., 2013", "Extraction automatique d'information depuis CV", "Q2", "ECIS"],
        ["[6]", "Reimers & Gurevych, 2019", "Sentence-BERT pour embeddings semantiques", "Q1", "EMNLP-IJCNLP"],
    ], [1.3, 3.4, 6.1, 1.5, 3])
    caption(doc, "Tableau 2 - Synthese globale des sources")
    h(doc, "2.3 Analyse critique", 2)
    p(doc, "Les travaux etudies confirment la pertinence des architectures content-based pour le demarrage a froid et la robustesse de TF-IDF comme baseline interpretable. Toutefois, ils couvrent rarement les specificites du marche marocain, l'integration complete parsing-matching-coaching et l'explicabilite granulaire des recommandations.")


def add_needs(doc):
    h(doc, "Chapitre 3 - Problematique, besoins et cahier des charges", 1)
    h(doc, "3.1 Problematique", 2)
    p(doc, "Comment concevoir une plateforme capable de recommander automatiquement des stages et emplois pertinents a des etudiants, tout en restant explicable, robuste, securisee, evolutive et adaptee au marche marocain ?")
    h(doc, "3.2 Acteurs du systeme", 2)
    table(doc, ["Acteur", "Responsabilites"], [
        ["Etudiant", "Profil, CV, recommandations, candidatures, coaching IA"],
        ["Entreprise", "Publication d'offres, suivi des candidatures, decision RH"],
        ["Administrateur", "Supervision, scraping, donnees, statistiques"],
    ], [3, 12])
    caption(doc, "Tableau 3 - Acteurs et responsabilites")
    h(doc, "3.3 Objectifs fonctionnels", 2)
    for item in [
        "Authentification securisee et gestion multi-roles.",
        "Upload, parsing et enrichissement du profil etudiant.",
        "Publication et gestion des offres par les entreprises.",
        "Candidature en ligne et suivi des statuts.",
        "Recommandations personnalisees avec score explicable.",
        "Generation de lettre de motivation, entretien simule et tendances marche.",
    ]:
        bullet(doc, item)
    h(doc, "3.4 Objectifs techniques", 2)
    for item in [
        "API REST Django/DRF securisee par JWT.",
        "Frontend React + TypeScript modulaire.",
        "Moteur NLP hybride TF-IDF, heuristiques et Gemini optionnel.",
        "Scraping multi-sources et resilience par fallback.",
        "Base extensible vers PostgreSQL, Celery, Docker et CI/CD.",
    ]:
        bullet(doc, item)


def add_architecture(doc):
    h(doc, "Chapitre 4 - Architecture globale du systeme", 1)
    h(doc, "4.1 Vue d'ensemble", 2)
    figure_flow(doc, "Figure 2 - Architecture logique de la plateforme", ["React + TypeScript\nSPA", "Axios + JWT\nAPI calls", "Django REST\nUsers / Offers / AI", "SQLite dev\nPostgreSQL cible", "NLP + Scraping\nServices Python"])
    table(doc, ["Couche", "Technologie", "Role"], [
        ["Frontend", "React 18, TypeScript, Vite", "Interface utilisateur, parcours etudiant/entreprise"],
        ["Backend", "Django, Django REST Framework", "API, authentification, permissions, logique metier"],
        ["Donnees", "SQLite puis PostgreSQL", "Persistance utilisateurs, offres, candidatures, recommandations"],
        ["IA/NLP", "pdfplumber, scikit-learn, Gemini API", "Extraction CV, matching, insights, generation"],
        ["Scraping", "requests, BeautifulSoup", "Collecte et normalisation d'offres"],
    ], [3, 5, 7])
    caption(doc, "Tableau 4 - Couches techniques du systeme")
    h(doc, "4.2 Flux principal", 2)
    figure_flow(doc, "Figure 3 - Pipeline fonctionnel de recommandation", ["Upload CV", "Extraction texte", "Parsing profil", "Vectorisation TF-IDF", "Matching offres", "Score + insights", "Candidature"])
    h(doc, "4.3 Justification des choix", 2)
    p(doc, "Le decoupage frontend/backend favorise la maintenance et l'evolution vers des clients mobiles. Django REST Framework apporte une base solide pour l'API et les permissions, tandis que React permet une experience interactive. TF-IDF est retenu comme baseline interpretable, economique et robuste pour une premiere version.")


def add_backend(doc):
    h(doc, "Chapitre 5 - Explication du code backend Django", 1)
    h(doc, "5.1 Organisation generale", 2)
    table(doc, ["Module", "Fichiers principaux", "Responsabilite"], [
        ["core", "settings.py, urls.py", "Configuration Django, apps installees, routes racines, CORS, JWT"],
        ["users", "models.py, serializers.py, views.py, urls.py", "User, StudentProfile, CompanyProfile, authentification, upload CV"],
        ["offers", "models.py, serializers.py, views.py, scraper.py", "Offres, candidatures, permissions metier, scraping"],
        ["ai", "models.py, services.py, views.py", "Recommandations, insights, traitement NLP, generation IA"],
    ], [2.7, 5, 7.3])
    caption(doc, "Tableau 5 - Modules backend et roles")
    h(doc, "5.2 Models users", 2)
    p(doc, "Le fichier backend/users/models.py definit un modele User heritant de AbstractUser et ajoute le champ role. Les proprietes is_student et is_company simplifient les controles metier. StudentProfile conserve les donnees academiques, le CV, le texte extrait, les competences detectees et les preferences. CompanyProfile stocke les informations de l'entreprise.")
    h(doc, "5.3 Models offers", 2)
    p(doc, "InternshipOffer represente une offre avec titre, description, competences requises, localisation, remuneration, statut et metadonnees de scraping. Application relie un etudiant a une offre et impose une contrainte unique afin d'eviter les candidatures dupliquees.")
    h(doc, "5.4 Serializers et views", 2)
    p(doc, "Les serializers transforment les objets Django en JSON et valident les entrees. Les views exposent les endpoints REST : inscription, connexion, profil, upload CV, liste des offres, creation d'offre, candidature, changement de statut et recuperation des recommandations.")
    h(doc, "5.5 Permissions", 2)
    p(doc, "La couche permission garantit qu'un etudiant ne peut postuler qu'en tant que candidat, qu'une entreprise ne modifie que ses propres offres, et que les candidatures recues restent visibles uniquement par l'entreprise proprietaire.")
    h(doc, "5.6 Endpoints principaux", 2)
    table(doc, ["Endpoint", "Role", "Description"], [
        ["POST /api/users/register/", "Public", "Creation d'un compte student ou company"],
        ["POST /api/users/token/", "Public", "Connexion JWT"],
        ["POST /api/users/token/refresh/", "Public", "Renouvellement du token"],
        ["GET/PUT /api/users/me/", "Connecte", "Lecture et mise a jour du profil"],
        ["POST /api/users/me/cv/", "Etudiant", "Upload CV et declenchement NLP"],
        ["GET/POST /api/offers/", "Public/Entreprise", "Lister ou publier une offre"],
        ["POST /api/offers/{id}/apply/", "Etudiant", "Candidater a une offre"],
        ["GET /api/ai/recommendations/", "Etudiant", "Top recommandations personnalisees"],
    ], [5, 2.5, 7.5])
    caption(doc, "Tableau 6 - API REST principale")


def add_frontend(doc):
    h(doc, "Chapitre 6 - Explication du code frontend React + TypeScript", 1)
    h(doc, "6.1 Structure du frontend", 2)
    table(doc, ["Dossier/Fichier", "Role"], [
        ["src/pages/", "Pages metier : Home, Login, Register, Profile, Offers, Recommendations, MyApplications, MyOffers, MarketTrends"],
        ["src/components/", "Composants reutilisables : Navbar, Layout, ProfileUpload, ExtractedProfile, CareerInsights, CoverLetterModal, InterviewModal"],
        ["src/contexts/AuthContext.tsx", "Etat global d'authentification, utilisateur courant, login/logout"],
        ["src/services/api.ts", "Client Axios centralise, injection JWT, refresh automatique en cas de 401"],
        ["src/App.tsx", "Routage principal et protection des routes"],
    ], [4.5, 10.5])
    caption(doc, "Tableau 7 - Organisation du frontend")
    h(doc, "6.2 Client API Axios", 2)
    p(doc, "Le fichier frontend/src/services/api.ts definit l'URL API, ajoute automatiquement Authorization: Bearer <token> dans chaque requete et intercepte les erreurs 401. Si le refresh_token existe, il appelle /users/token/refresh/, remplace l'access token et rejoue la requete initiale.")
    h(doc, "6.3 Parcours etudiant", 2)
    for item in [
        "Creation de compte puis connexion.",
        "Completion du profil et upload du CV.",
        "Affichage du profil extrait et des competences detectees.",
        "Consultation des recommandations avec score, competences alignees et manquantes.",
        "Generation d'une lettre de motivation et simulation d'entretien.",
        "Depot et suivi des candidatures.",
    ]:
        numbered(doc, item)
    h(doc, "6.4 Parcours entreprise", 2)
    for item in [
        "Creation d'un compte entreprise.",
        "Publication d'offres avec competences requises.",
        "Visualisation des candidatures recues.",
        "Changement du statut : pending, accepted, rejected ou withdrawn.",
    ]:
        numbered(doc, item)


def add_ai_scraping(doc):
    h(doc, "Chapitre 7 - Explication du code IA, NLP et scraping", 1)
    h(doc, "7.1 Service NLP", 2)
    p(doc, "Le fichier backend/ai/services.py contient la classe NLPService. Elle normalise le texte, canonise les competences, extrait les textes de CV PDF/DOCX, detecte les competences, estime le niveau d'experience et calcule les recommandations.")
    table(doc, ["Fonction", "Explication"], [
        ["normalize_text", "Nettoie et standardise le texte pour limiter les variations inutiles."],
        ["extract_text_from_cv", "Lit les PDF avec pdfplumber et les DOCX avec python-docx."],
        ["extract_profile_with_ai", "Utilise Gemini si la cle API existe, sinon fallback heuristique."],
        ["compute_similarity", "Construit une matrice TF-IDF puis calcule la similarite cosinus."],
        ["compute_skill_overlap", "Mesure le chevauchement entre competences du CV et competences de l'offre."],
        ["compute_context_bonus", "Ajoute des bonus contexte : stage, Maroc, Junior, remuneration."],
        ["process_student_cv", "Pipeline complet : extraction, enrichissement profil, scoring et sauvegarde des recommandations."],
        ["build_career_insights", "Agrege les recommandations pour produire forces, gaps et plan d'action."],
    ], [4.3, 10.7])
    caption(doc, "Tableau 8 - Fonctions principales du moteur NLP")
    h(doc, "7.2 Formule de scoring", 2)
    add_mixed(doc, [{"text": "Score final = 0,65 x similarite semantique + 0,30 x overlap competences + bonus contexte/preference", "bold": True, "color": GREEN}], align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Cette formule combine un signal lexical global, un signal metier sur les competences et un signal de personnalisation. Le resultat est converti en score interpretable pour l'utilisateur.")
    h(doc, "7.3 Explainable AI", 2)
    bullet(doc, "semantic_score : proximite globale entre CV et offre.")
    bullet(doc, "skill_overlap_score : pourcentage de competences communes.")
    bullet(doc, "matching_skills : competences deja presentes chez l'etudiant.")
    bullet(doc, "missing_skills : competences a renforcer.")
    bullet(doc, "score_band : excellent, fort, prometteur ou a explorer.")
    h(doc, "7.4 Scraping et intelligence marche", 2)
    p(doc, "Le module backend/offers/scraper.py collecte des offres externes, les nettoie, detecte les champs utiles, normalise les donnees et les insere ou met a jour dans la base. La commande de management permet de declencher cette collecte depuis l'administration ou une tache planifiee.")
    figure_flow(doc, "Figure 4 - Pipeline scraping et normalisation", ["Source web", "HTML", "Parsing", "Nettoyage", "Filtre Maroc", "Extraction skills", "Upsert DB"])


def add_security_validation(doc):
    h(doc, "Chapitre 8 - Securite, performance, validation et limites", 1)
    h(doc, "8.1 Securite", 2)
    table(doc, ["Aspect", "Implementation actuelle", "Renforcement production"], [
        ["Authentification", "JWT access + refresh", "Rotation, durees strictes, blacklist"],
        ["Autorisations", "Roles et proprietaires", "Tests automatiques de permissions"],
        ["Validation", "Serializers DRF", "Audit OWASP API Top 10"],
        ["CORS", "Origines autorisees", "Politique stricte par environnement"],
        ["Secrets", "Variables d'environnement", "Vault ou secret manager"],
    ], [3, 6, 6])
    caption(doc, "Tableau 9 - Synthese securite")
    h(doc, "8.2 Performance et scalabilite", 2)
    bullet(doc, "Deplacer le traitement NLP et le scraping vers Celery + Redis.")
    bullet(doc, "Mettre en cache les vecteurs TF-IDF des offres non modifiees.")
    bullet(doc, "Migrer SQLite vers PostgreSQL avec indexation adaptee.")
    bullet(doc, "Stocker les CV sur un stockage objet type S3 ou Cloudflare R2.")
    bullet(doc, "Ajouter logs structures, monitoring et tests de charge.")
    h(doc, "8.3 Validation qualitative", 2)
    table(doc, ["Scenario", "Resultat"], [
        ["Matching etudiant", "Recommandations coherentes lorsque CV et competences sont renseignes."],
        ["Parcours entreprise", "Creation d'offres et gestion des candidatures operationnelles."],
        ["Outils IA", "Lettre de motivation et entretien simule disponibles."],
        ["Scraping", "Collecte et fallback prevus pour assurer la continuite."],
        ["Securite", "Controle d'acces par role et propriete des ressources."],
    ], [4, 11])
    caption(doc, "Tableau 10 - Validation par scenario")
    h(doc, "8.4 Limites", 2)
    bullet(doc, "Couverture de tests encore insuffisante.")
    bullet(doc, "Matching principalement lexical en V1.")
    bullet(doc, "Pas encore de collaborative filtering comportemental.")
    bullet(doc, "Scalabilite non validee par tests de charge.")
    bullet(doc, "Domaine de scraping limite a un volume initial d'offres.")


def add_positioning(doc):
    h(doc, "Chapitre 9 - Positionnement, innovation et perspectives de recherche", 1)
    h(doc, "9.1 Lacunes identifiees", 2)
    table(doc, ["Lacune", "Contribution du projet"], [
        ["Marche marocain peu adresse", "Scraping et logique metier adaptes au contexte local."],
        ["Solutions souvent fragmentees", "Plateforme integree : parsing, matching, candidature, coaching."],
        ["Explicabilite limitee", "Scores decomposes, competences alignees/manquantes, resume naturel."],
        ["Peu de personnalisation etudiant", "Preferences de villes, types d'offres, intitules cibles et remote."],
    ], [5, 10])
    caption(doc, "Tableau 11 - Lacunes et axes d'innovation")
    h(doc, "9.2 Perspectives de recherche", 2)
    bullet(doc, "Sentence-BERT pour depasser les limites synonymiques de TF-IDF.")
    bullet(doc, "Evaluation formelle : Precision@K, Recall@K, nDCG, MRR.")
    bullet(doc, "Collaborative filtering a partir des clics, favoris et candidatures.")
    bullet(doc, "Detection de biais algorithmiques selon genre, region ou niveau social.")
    bullet(doc, "Support multilingue francais, arabe et darija.")
    bullet(doc, "Apprentissage federe pour proteger les donnees personnelles.")


def add_conclusion_refs(doc):
    h(doc, "Conclusion et roadmap", 1)
    p(doc, "Ce PFA a permis de concevoir et de mettre en oeuvre une plateforme intelligente de recommandation de stages et d'emplois. La solution combine une API Django REST securisee, un frontend React moderne, un moteur NLP interpretable, des modules IA generatifs et une veille marche par scraping.")
    p(doc, "La valeur du projet reside dans sa capacite a transformer la recherche d'opportunites en experience proactive : l'etudiant comprend pourquoi une offre lui est recommandee, l'entreprise centralise ses candidatures et l'administrateur pilote l'alimentation des donnees.")
    table(doc, ["Horizon", "Evolution"], [
        ["Court terme (3-6 mois)", "Tests automatises, PostgreSQL, dockerisation, CI/CD GitHub Actions."],
        ["Moyen terme (6-12 mois)", "Celery + Redis, cache TF-IDF, evaluation Precision@K/nDCG, dashboard analytique."],
        ["Long terme (12+ mois)", "Sentence-BERT, collaborative filtering, application mobile, conformite Loi 09-08/RGPD, modele SaaS."],
    ], [4, 11])
    caption(doc, "Tableau 12 - Roadmap d'evolution")
    doc.add_page_break()
    h(doc, "References bibliographiques", 1)
    refs = [
        '[1] F. O. Isinkaye, Y. O. Folajimi, and B. A. Ojokoh, "Recommendation systems: Principles, methods and evaluation," Egyptian Informatics Journal, vol. 16, no. 3, pp. 261-273, 2015.',
        '[2] K. Siting, W. Wenxing, Z. Ning, and Y. Fan, "Job recommender systems: A survey," Proc. 7th Int. Conf. Computer Science & Education (ICCSE), pp. 920-924, 2012.',
        '[3] T. Al-Otaibi and M. Ykhlef, "A survey of job recommender systems," International Journal of Physical Sciences, vol. 7, no. 29, pp. 5127-5142, 2012.',
        '[4] D. Yi, C. Wei, and J. Ruifeng, "A content-based job recommendation system using TF-IDF and cosine similarity," Proc. Int. Conf. Machine Learning and Computing, 2020.',
        '[5] E. Kops, G. Kessler, and T. Durschner, "Automated resume information extraction to support the matching of job seekers and vacancies," Proc. 21st European Conf. Information Systems (ECIS), 2013.',
        '[6] N. Reimers and I. Gurevych, "Sentence-BERT: Sentence embeddings using Siamese BERT-networks," Proc. EMNLP-IJCNLP, pp. 3982-3992, 2019.',
        '[7] J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, "BERT: Pre-training of deep bidirectional transformers for language understanding," Proc. NAACL-HLT, pp. 4171-4186, 2019.',
        '[8] T. Mikolov, I. Sutskever, K. Chen, G. Corrado, and J. Dean, "Distributed representations of words and phrases and their compositionality," Proc. NIPS, pp. 3111-3119, 2013.',
        '[9] Django Software Foundation, "Django REST Framework Documentation," Version 3.14, 2024. [Online]. Available: https://www.django-rest-framework.org/',
        '[10] F. Pedregosa et al., "Scikit-learn: Machine learning in Python," Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.',
        '[11] Google LLC, "Gemini API Documentation," 2024. [Online]. Available: https://ai.google.dev/docs',
        '[12] Meta Platforms, Inc., "React Documentation," Version 18, 2024. [Online]. Available: https://react.dev/',
    ]
    for ref in refs:
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(4)
        run = par.add_run(ref)
        run.font.name = "Arial"
        run.font.size = Pt(9)


def build():
    doc = Document()
    style_doc(doc)
    header_footer(doc)
    add_cover(doc)
    h(doc, "Table des matieres", 1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()
    add_intro(doc)
    doc.add_page_break()
    add_conceptual(doc)
    doc.add_page_break()
    add_methodology(doc)
    doc.add_page_break()
    add_needs(doc)
    doc.add_page_break()
    add_architecture(doc)
    doc.add_page_break()
    add_backend(doc)
    doc.add_page_break()
    add_frontend(doc)
    doc.add_page_break()
    add_ai_scraping(doc)
    doc.add_page_break()
    add_security_validation(doc)
    doc.add_page_break()
    add_positioning(doc)
    doc.add_page_break()
    add_conclusion_refs(doc)
    mark_update_fields(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
